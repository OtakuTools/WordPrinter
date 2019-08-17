from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json, time, lxml, threading
import re, os
from getTime import getTime
from presetData import getColorStyle,getParagraphStyle
from dataStruct import List2Dot


class Replace:
    lock = threading.RLock()

    def __init__(self):
        self.timeCount = 0

    def __del__(self):
        pass
   
    def rule(self, r):
        try:
            #取颜色对应的属性
            obj = self
            for attrName in self.colorDict[str(r.font.color.rgb)].split('.'):
                obj = getattr( obj , attrName, r.text )
            
            #针对modifyDate
            if str(r.font.color.rgb) == 'F50000':
                obj = getattr(obj,'index'+str(self.timeCount%5+1))
                self.timeCount += 1
            
            #针对多行文本
            if type(obj) == type([]):
                obj = '\n'.join(obj)

            r.text = str(obj)
            r.font.highlight_color = None
            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        except Exception as e:
            print(str(e))

    def findAndReplaceFootAndHead(self):
        #页眉页脚
        self.lock.acquire()
        try:
            for s in self.document.sections:
                for p in s.footer.paragraphs:
                    for r in p.runs:
                        if r.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                            r.text = self.user.company
                            r.font.highlight_color = None
                            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                for p in s.header.paragraphs:
                    for r in p.runs:
                        if r.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                            r.text = self.user.fileName
                            r.font.highlight_color = None
                            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        except Exception as e:
            print(e)
        finally:
            self.lock.release()

    def findAndReplaceTables(self):
        #表格
        for t in self.document.tables:
            for c in t.columns:
                for cc in c.cells:
                    for p in cc.paragraphs:
                        for r in p.runs:
                            if r.font.highlight_color == WD_COLOR_INDEX.YELLOW :
                                try:
                                    self.lock.acquire()
                                    #self.replaceRules(r)
                                    self.rule(r)
                                except Exception as e:
                                    print('表格'+str(e))
                                    raise
                                finally:
                                    self.lock.release()

    def replaceDepartmentTable(self):
        #服务管理职责分配表
        try:
            table = self.document.tables[-1]
            table_row_len = len(table.rows)
            #s = time.time()
            for d in self.user.departments:
                #特殊情况
                if d["name"] == "管理者代表":
                    continue
                # 表头
                column = table.add_column( Cm(3) )
                cell = column.cells[0]
                #特别情况
                if d["name"] == "总经理":
                    cell.text = "管理层"
                else:
                    cell.text = d["name"]
                cell.paragraphs[0].style = "表格标记"
                # 内容
                # 优化替换速度，用set取代list，平均速度提高可达10%
                funcSet = set(d["func"])
                for i in range(1, table_row_len):
                    cell = column.cells[i]
                    if i in funcSet:
                        cell.text = "▲"
                    else:
                        cell.text = "△"
                    cell.paragraphs[0].style = "表格标记"
            table.style = 'Table Theme'
            table.autofit = True
        except Exception as e:
            raise

    def findAndReplaceParagraphs(self):
        #正文
        for p in self.document.paragraphs:
            for r in p.runs:
                if r.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    # 部门介绍
                    if str(r.font.color.rgb) == 'EE0000':
                        p.clear()
                        for d in self.user.departments:
                            p.insert_paragraph_before( d['name'] + ':' , '部门名称' )
                            for i in d['intro']:
                                p.insert_paragraph_before( i , '部门职责' )
                    '''
                    # 公司简介
                    if str(r.font.color.rgb) == 'FD0000':
                        p.clear()
                        for intro in self.user.introduction:
                            p.insert_paragraph_before( intro , '简介' )
                    '''

                    # 段落替换
                    paragraphDict = getParagraphStyle()
                    if str(r.font.color.rgb) in paragraphDict:
                        #get attr
                        obj = self
                        for attrName in paragraphDict[str(r.font.color.rgb)]['attr'].split('.'):
                            obj = getattr( obj , attrName, [] )
                        #replace
                        p.clear()
                        for paragraph in obj:
                            p.insert_paragraph_before( paragraph ,  paragraphDict[str(r.font.color.rgb)]['style'] )

                    # 插入图片
                    if str(r.font.color.rgb) == 'ED0000':
                        pp = p.insert_paragraph_before()
                        #检查宽高
                        levelDict = {}
                        for dep in self.user.departments:
                            if not levelDict.__contains__(dep['level']):
                                levelDict[dep['level']] = 1
                            else:
                                levelDict[dep['level']] = levelDict[dep['level']] + 1
                        height = len( levelDict )
                        width = max(levelDict.values())
                        pp.add_run().add_picture( self.user.picPath ,height=Cm(10)) if width < 2*height else pp.add_run().add_picture( self.user.picPath ,width=Cm(16))
                        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # 插入logo
                    if str(r.font.color.rgb) == 'EF0000':
                        #insert logo
                        if(self.user.logoPath != '' and os.path.exists(self.user.logoPath)):
                            pp = p.insert_paragraph_before()
                            pp.add_run().add_picture( self.user.logoPath , height=Cm(4.5) )
                            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        #clear text
                        p.clear()

                    # 部门名称
                    if str(r.font.color.rgb) == 'D50000':
                        r.text = ''
                        for dep in self.user.departments:
                            if dep['name'] != '管理层' and dep['name'] != '总经理' and dep['name'] != '管理者代表':
                                r.text += ( '、'+dep['name'] )
                        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                        r.font.highlight_color = None

                    try:
                        self.lock.acquire()
                        #self.replaceRules(r)
                        self.rule(r)
                    except Exception as e:
                        raise
                    finally:
                        self.lock.release()

    def run(self, src , dst , user, project):
        self.document = Document(src)
        self.user = getTime(user)
        self.colorDict = getColorStyle()
        #针对modifyDate
        dateList = self.user.modifyDate
        self.user.modifyDate = List2Dot(self.user.modifyDate)

        #获取项目对象
        self.project = project
        if self.project != None:
            event = self.project.ServiceProcess.Event
            #受理事件数
            self.accepted = event.S1 + event.S2 + event.S3 + event.S4
            #重大事件比率
            self.criticalRate = event.S1 / self.accepted
            #事件关闭率
            self.closeRate = self.accepted / event.closed
        
        #更新目录
        element_updatefields = lxml.etree.SubElement(
            self.document.settings.element, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"+"updateFields"
        )
        element_updatefields.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"+"val", "true")

        threadList = []
        threadList.append(threading.Thread(target=self.findAndReplaceFootAndHead))
        threadList.append(threading.Thread(target=self.findAndReplaceTables))
        threadList.append(threading.Thread(target=self.findAndReplaceParagraphs))

        #一层SM文件有职责表
        if re.search( "(.*)-SM-(.*)" , src ):
            self.replaceDepartmentTable()

        for t in threadList:
            t.start()
        for t in threadList:
            t.join()
            
        print('成功生成 '+dst )

        #针对modifyDate
        self.user.modifyDate = dateList

        return self.document
