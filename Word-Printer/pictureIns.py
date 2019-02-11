from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Cm

def pictureIns(src,dst,pic):
    document = Document(src)
    for p in document.paragraphs:
        for r in p.runs:
            if r.font.highlight_color == WD_COLOR_INDEX.YELLOW and str(r.font.color.rgb) == '000FFF':
                r.font.highlight_color = None
                r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                pp = p.insert_paragraph_before()
                pp.add_run().add_picture(pic,Cm(16))
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(dst)
    print('图片插入成功')
    pass
